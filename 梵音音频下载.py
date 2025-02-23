import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import time
import requests
from docx import Document
import fitz  # PyMuPDF
import re
from gtts import gTTS
from typing import *
import logging
from PIL import Image, ImageTk
import json
import winreg
import urllib.request
import socket


def get_system_proxy():
    """获取系统代理设置"""
    try:
        reg_path = r'Software\Microsoft\Windows\CurrentVersion\Internet Settings'
        reg_key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, reg_path, 0, winreg.KEY_ALL_ACCESS)
        proxy_enable, _ = winreg.QueryValueEx(reg_key, 'ProxyEnable')
        if proxy_enable:
            proxy_server, _ = winreg.QueryValueEx(reg_key, 'ProxyServer')
            if proxy_server:
                if ':' in proxy_server:
                    host, port = proxy_server.split(':')
                    return host.strip(), port.strip()
        http_proxy = os.environ.get('HTTP_PROXY') or os.environ.get('http_proxy')
        https_proxy = os.environ.get('HTTPS_PROXY') or os.environ.get('https_proxy')
        if http_proxy:
            proxy = http_proxy
        elif https_proxy:
            proxy = https_proxy
        else:
            return None, None
        if proxy.startswith('http://'):
            proxy = proxy[7:]
        elif proxy.startswith('https://'):
            proxy = proxy[8:]
        if '@' in proxy:
            proxy = proxy.split('@')[1]
        if ':' in proxy:
            host, port = proxy.split(':')
            return host.strip(), port.strip()
    except Exception as e:
        logging.warning(f"获取系统代理设置失败: {str(e)}")
    return None, None


def batch_text_to_speech(file_path, folder_path, progress_callback=None):
    """批量转换Word和PDF文件中的纯梵文段落为音频"""
    input_filename = os.path.splitext(os.path.basename(file_path))[0]
    folder_path = os.path.join(folder_path, input_filename)

    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    texts = []
    titles = []

    total_progress = 0
    current_progress = 0

    try:
        if file_path.lower().endswith('.pdf'):
            pdf_document = fitz.open(file_path)
            total_progress = len(pdf_document)
            logging.info(f"开始处理PDF文件: {file_path}")

            for page_index, page in enumerate(pdf_document):
                # Get detailed font information
                page_text_blocks = page.get_text("dict")["blocks"] # Get dict for more details
                logging.info(f"处理第 {page_index + 1} 页")

                current_title = None
                text_to_convert = ""

                for block_index, block in enumerate(page_text_blocks):
                    logging.debug(f"  Block {block_index}: {block}")  # 打印整个 block
                    if "lines" in block:
                        for line_index, line in enumerate(block["lines"]):
                            logging.debug(f"    Line {line_index}: {line}") # 打印整个 line
                            for span_index, span in enumerate(line["spans"]):
                                font_name = span["font"]
                                # font_size = span["size"] # Size is less critical now
                                text = span["text"].strip()
                                logging.info(f"      Span {span_index}: Font={font_name}, Text='{text}'") # 关键：打印字体和文本
                                logging.debug(f"        current_title: {current_title}, text_to_convert: {text_to_convert}") #打印状态

                                if "YaHei" in font_name :
                                    if text_to_convert and current_title:  # Previous mantra is complete
                                        logging.info(f"        Adding to texts: title='{current_title}', text='{text_to_convert}'")
                                        texts.append(text_to_convert.strip())
                                        titles.append(current_title.strip())
                                    # 提取标题 (处理编号和卍字符)
                                    match = re.match(r"^(M\d+\.\d+)?\s*(卍\s*)*(.*)$", text)
                                    if match:
                                        title_text = match.group(3).strip()  # 提取标题文本
                                        if title_text:  # 仅当提取到有效标题时才更新
                                            current_title = title_text
                                    text_to_convert = "" #Reset
                                elif (font_name.startswith("Arial") or font_name.startswith("Times")) and current_title:
                                    # 罗马音译咒语, 且已有标题
                                    text_to_convert += text + " "  # 累加咒语文本, 用空格分隔

                    else:
                        logging.debug("    Block has no lines.")

                #Check remaining
                if text_to_convert and current_title:
                    logging.info(f"  Adding remaining text: title='{current_title}', text='{text_to_convert}'")
                    texts.append(text_to_convert.strip())
                    titles.append(current_title.strip())

                current_progress += 1
                if progress_callback:
                    progress_callback(current_progress, total_progress)

            pdf_document.close()


        elif file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            total_progress = len(doc.paragraphs)
            logging.info(f"开始处理Word文件: {file_path}")


            current_title = None
            text_to_convert = ""

            for para in doc.paragraphs:
                for run in para.runs:
                    font_name = run.font.name
                    text = run.text.strip()

                    if "YaHei" in font_name : # 标题
                        if text_to_convert and current_title:
                            texts.append(text_to_convert.strip())
                            titles.append(current_title.strip())
                            text_to_convert = ""

                        match = re.match(r"^(M\d+\.\d+)?\s*(卍\s*)*(.*)$", text)
                        if match:
                            title_text = match.group(3).strip()
                            if title_text:
                                current_title = title_text
                    elif (font_name and font_name.startswith("Arial") or font_name and font_name.startswith("Times")) and current_title :
                         text_to_convert += text + " "

            if text_to_convert and current_title:
                texts.append(text_to_convert.strip())
                titles.append(current_title.strip())

            current_progress = 0 # reset
            for para_index, para in enumerate(doc.paragraphs):
                current_progress += 1
                if progress_callback:
                    progress_callback(current_progress, total_progress)

        else:
            error_msg = "不支持的文件类型，请提供 .docx 或 .pdf 文件"
            logging.error(error_msg)
            raise ValueError(error_msg)

        audio_paths = []
        total_progress = len(texts)
        current_progress = 0

        record_path = os.path.join(folder_path, "audio_record.txt")

        for text, title in zip(texts, titles):
            try:
                if not text.strip():
                    logging.warning(f"跳过空文本，标题: {title}")
                    continue

                audio_name = title.replace("/", "-").replace("\\", "-")  # Sanitize filename
                audio_path = os.path.join(folder_path, f"{audio_name}.mp3")

                with open(record_path, "a", encoding='utf-8') as f:
                    f.write(f"音频文件: {audio_name}.mp3\n标题: {title}\n内容: {text}\n{'='*50}\n")

                logging.info(f"开始转换音频: {audio_name}.mp3")
                logging.info(f"文本长度: {len(text)} 字符")
                convert_to_audio(text, audio_path, lang='ro') # Romanian!

                if os.path.exists(audio_path) and os.path.getsize(audio_path) > 0:
                    audio_paths.append(audio_path)
                    logging.info(f"成功转换: {audio_name}.mp3 (大小: {os.path.getsize(audio_path)} 字节)")

                else:
                    logging.error(f"音频文件生成失败或为空: {audio_name}.mp3")

                current_progress += 1
                if progress_callback:
                    progress_callback(current_progress, total_progress)

            except Exception as e:
                logging.error(f"转换失败 '{title}': {str(e)}")
                logging.error(f"错误详情: {e.__class__.__name__}")
                continue

        return audio_paths

    except Exception as e:
        logging.error(f"处理文件时发生错误: {str(e)}")
        raise
    finally:
        pass

def convert_to_audio(text, audio_path, lang='ro', max_retries=3, retry_delay=2):
    """转换文本为音频，支持指定语言"""
    for attempt in range(max_retries):
        try:
            logging.info(f"尝试生成音频文件 {audio_path} (第 {attempt + 1} 次尝试)")
            logging.debug(f"文本内容: {text[:100]}...")

            if not text.strip():
                raise ValueError("文本内容为空")

            try:
                socket.create_connection(("translate.google.com", 80), timeout=10)
            except (requests.exceptions.RequestException, socket.error) as e:
                raise requests.exceptions.RequestException(f"无法连接到Google服务: {str(e)}")

            tts = gTTS(text=text, lang=lang)  # Use specified language

            os.makedirs(os.path.dirname(audio_path), exist_ok=True)
            tts.save(audio_path)

            if not os.path.exists(audio_path):
                raise Exception("音频文件未生成")

            file_size = os.path.getsize(audio_path)
            if file_size == 0:
                os.remove(audio_path)
                raise Exception("生成的音频文件大小为0")

            logging.info(f"成功生成音频文件: {audio_path} (大小: {file_size} 字节)")
            return

        except requests.exceptions.RequestException as e:
            error_msg = f"网络请求错误 (第 {attempt + 1} 次尝试): {str(e)}"
            logging.error(error_msg)
            logging.debug(f"网络错误详情: {e.__class__.__name__}")

            if attempt == max_retries - 1:
                raise Exception(f"音频生成失败: {error_msg}")

            retry_delay_adjusted = retry_delay * (2 ** attempt)
            logging.info(f"等待 {retry_delay_adjusted} 秒后重试...")
            time.sleep(retry_delay_adjusted)

        except Exception as e:
            error_msg = f"音频生成错误 (第 {attempt + 1} 次尝试): {str(e)}"
            logging.error(error_msg)
            logging.debug(f"错误类型: {e.__class__.__name__}, 错误详情: {str(e)}")

            if os.path.exists(audio_path):
                try:
                    os.remove(audio_path)
                except Exception as remove_error:
                    logging.warning(f"无法删除损坏的音频文件: {str(remove_error)}")

            if attempt == max_retries - 1:
                raise Exception(f"音频生成失败: {error_msg}")

            retry_delay_adjusted = retry_delay * (2 ** attempt)
            logging.info(f"等待 {retry_delay_adjusted} 秒后重试...")
            time.sleep(retry_delay_adjusted)



class SanskritAudioConverterGUI(tk.Tk):
    def __init__(self):
        super().__init__()

        self.title("梵文音频批量转换工具 v1.6")  # Updated version
        self.geometry("800x800")
        try:
            self.iconbitmap("app_icon.ico")
        except:
            pass
        self.config_file = "audio_converter_config.json"
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        self.create_file_selection_frame(main_frame)
        self.create_proxy_frame(main_frame)
        self.create_text_input_frame(main_frame)
        self.create_progress_frame(main_frame)
        self.toggle_proxy_fields()
        self.load_config()
        self.check_single_instance()

    def create_text_input_frame(self, main_frame):
        """创建文本输入和播放框架"""
        text_frame = ttk.LabelFrame(main_frame, text="文本播放", padding="5")
        text_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.text_input = tk.Text(text_frame, height=10)
        self.text_input.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        lang_frame = ttk.Frame(text_frame)
        lang_frame.pack(fill=tk.X, padx=5, pady=2)
        ttk.Label(lang_frame, text="语言:").pack(side=tk.LEFT)
        self.lang_var = tk.StringVar(value="ro")  # Default to Romanian
        lang_combo = ttk.Combobox(lang_frame, textvariable=self.lang_var, values=["sa", "hi", "en", "ro", "id"])
        lang_combo.pack(side=tk.LEFT, padx=5)
        button_frame = ttk.Frame(text_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(button_frame, text="播放", command=self.play_text).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="下载", command=self.download_text).pack(side=tk.LEFT, padx=5)

    def play_text(self):
        """播放输入的文本"""
        text = self.text_input.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("警告", "请输入要播放的文本")
            return
        try:
            proxy = self.get_current_proxy()
            if proxy:
                os.environ['HTTP_PROXY'] = proxy['http']
                os.environ['HTTPS_PROXY'] = proxy['https']
            temp_dir = "temp"
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)
            temp_audio = os.path.join(temp_dir, "temp_audio.mp3")
            tts = gTTS(text=text, lang=self.lang_var.get())
            tts.save(temp_audio)
            os.startfile(temp_audio)
        except Exception as e:
            messagebox.showerror("错误", f"播放失败: {str(e)}")
        finally:
            os.environ.pop('HTTP_PROXY', None)
            os.environ.pop('HTTPS_PROXY', None)

    def download_text(self):
        """下载文本的音频文件"""
        text = self.text_input.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("警告", "请输入要转换的文本")
            return
        try:
            proxy = self.get_current_proxy()
            if proxy:
                os.environ['HTTP_PROXY'] = proxy['http']
                os.environ['HTTPS_PROXY'] = proxy['https']
            file_path = filedialog.asksaveasfilename(
                defaultextension=".mp3",
                filetypes=[("MP3文件", "*.mp3")]
            )
            if not file_path:
                return
            tts = gTTS(text=text, lang=self.lang_var.get())
            tts.save(file_path)
            messagebox.showinfo("成功", "音频文件已保存")
        except Exception as e:
            messagebox.showerror("错误", f"下载失败: {str(e)}")
        finally:
            os.environ.pop('HTTP_PROXY', None)
            os.environ.pop('HTTPS_PROXY', None)

    def create_proxy_frame(self, main_frame):
        """创建代理设置框架"""
        proxy_frame = ttk.LabelFrame(main_frame, text="代理设置", padding="5")
        proxy_frame.pack(fill=tk.X, padx=5, pady=5)
        self.auto_proxy_var = tk.BooleanVar(value=True)
        auto_proxy_check = ttk.Checkbutton(proxy_frame, text="自动获取系统代理",
                                          variable=self.auto_proxy_var,
                                          command=self.toggle_proxy_fields)
        auto_proxy_check.pack(fill=tk.X, pady=2)
        proxy_host_frame = ttk.Frame(proxy_frame)
        proxy_host_frame.pack(fill=tk.X, pady=2)
        proxy_host_label = ttk.Label(proxy_host_frame, text="代理地址:")
        proxy_host_label.pack(side=tk.LEFT)
        self.proxy_host_var = tk.StringVar(value="127.0.0.1")
        self.proxy_host_entry = ttk.Entry(proxy_host_frame, textvariable=self.proxy_host_var)
        self.proxy_host_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        proxy_port_frame = ttk.Frame(proxy_frame)
        proxy_port_frame.pack(fill=tk.X, pady=2)
        proxy_port_label = ttk.Label(proxy_port_frame, text="代理端口:")
        proxy_port_label.pack(side=tk.LEFT)
        self.proxy_port_var = tk.StringVar(value="7890")
        self.proxy_port_entry = ttk.Entry(proxy_port_frame, textvariable=self.proxy_port_var)
        self.proxy_port_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

    def load_config(self):
        """加载配置文件"""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    self.auto_proxy_var.set(config.get('auto_proxy', True))
                    self.proxy_host_var.set(config.get('proxy_host', '127.0.0.1'))
                    self.proxy_port_var.set(config.get('proxy_port', '7890'))
        except Exception as e:
            logging.warning(f"加载配置文件失败: {str(e)}")

    def save_config(self):
        """保存配置到文件"""
        config = {
            'auto_proxy': self.auto_proxy_var.get(),
            'proxy_host': self.proxy_host_var.get(),
            'proxy_port': self.proxy_port_var.get()
        }
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=4)
        except Exception as e:
            logging.warning(f"保存配置文件失败: {str(e)}")

    def check_single_instance(self):
        """确保只有一个程序实例在运行"""
        self.lock_file = "audio_converter.lock"
        try:
            if os.path.exists(self.lock_file):
                try:
                    with open(self.lock_file, 'r') as f:
                        pid = int(f.read().strip())
                    os.kill(pid, 0)
                    messagebox.showerror("错误", "程序已经在运行中")
                    self.quit()
                    return False
                except (ProcessLookupError, ValueError):
                    try:
                        os.remove(self.lock_file)
                    except:
                        pass
            with open(self.lock_file, 'w') as f:
                f.write(str(os.getpid()))
            return True
        except Exception as e:
            logging.error(f"检查单例模式时发生错误: {str(e)}")
            return False

    def __del__(self):
        """析构函数，确保清理锁文件"""
        try:
            if hasattr(self, 'lock_file') and os.path.exists(self.lock_file):
                os.remove(self.lock_file)
        except:
            pass

    def create_file_selection_frame(self, main_frame):
        """创建文件选择框架"""
        file_frame = ttk.LabelFrame(main_frame, text="文件选择", padding="5")
        file_frame.pack(fill=tk.X, padx=5, pady=5)
        input_frame = ttk.Frame(file_frame)
        input_frame.pack(fill=tk.X, pady=2)
        self.input_path_var = tk.StringVar()
        input_label = ttk.Label(input_frame, text="输入文件:")
        input_label.pack(side=tk.LEFT)
        input_entry = ttk.Entry(input_frame, textvariable=self.input_path_var)
        input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        input_button = ttk.Button(input_frame, text="浏览", command=self.select_input_file)
        input_button.pack(side=tk.LEFT)
        output_frame = ttk.Frame(file_frame)
        output_frame.pack(fill=tk.X, pady=2)
        self.output_path_var = tk.StringVar()
        output_label = ttk.Label(output_frame, text="输出目录:")
        output_label.pack(side=tk.LEFT)
        output_entry = ttk.Entry(output_frame, textvariable=self.output_path_var)
        output_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        output_button = ttk.Button(output_frame, text="浏览", command=self.select_output_folder)
        output_button.pack(side=tk.LEFT)

    def create_progress_frame(self, main_frame):
        """创建进度显示框架"""
        progress_frame = ttk.LabelFrame(main_frame, text="转换进度", padding="5")
        progress_frame.pack(fill=tk.X, padx=5, pady=5)
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(progress_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, padx=5, pady=5)
        self.status_var = tk.StringVar(value="就绪")
        status_label = ttk.Label(progress_frame, textvariable=self.status_var)
        status_label.pack(pady=2)
        self.convert_button = ttk.Button(progress_frame, text="开始转换", command=self.start_conversion)
        self.convert_button.pack(pady=5)

    def select_input_file(self):
        """选择输入文件"""
        file_types = [("支持的文件", "*.pdf;*.docx"), ("PDF文件", "*.pdf"), ("Word文件", "*.docx")]
        file_path = filedialog.askopenfilename(filetypes=file_types)
        if file_path:
            self.input_path_var.set(file_path)

    def select_output_folder(self):
        """选择输出目录"""
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.output_path_var.set(folder_path)

    def update_progress(self, current, total):
        """更新进度条"""
        progress = (current / total * 100) if total > 0 else 0
        self.progress_var.set(progress)
        self.status_var.set(f"处理中... {current}/{total}")
        self.update()

    def toggle_proxy_fields(self):
        """切换代理字段的可用状态"""
        state = 'disabled' if self.auto_proxy_var.get() else 'normal'
        self.proxy_host_entry.configure(state=state)
        self.proxy_port_entry.configure(state=state)
        if self.auto_proxy_var.get():
            self.update_proxy_settings()

    def update_proxy_settings(self):
        """更新代理设置"""
        if self.auto_proxy_var.get():
            host, port = get_system_proxy()
            if host and port:
                self.proxy_host_var.set(host)
                self.proxy_port_var.set(port)
            else:
                self.proxy_host_var.set("127.0.0.1")
                self.proxy_port_var.set("7890")

    def get_current_proxy(self):
        """获取当前代理设置"""
        if self.auto_proxy_var.get():
            host, port = get_system_proxy()
            if host and port:
                return {
                    'http': f'http://{host}:{port}',
                    'https': f'http://{host}:{port}'
                }
        host = self.proxy_host_var.get()
        port = self.proxy_port_var.get()
        if host and port:
            return {
                'http': f'http://{host}:{port}',
                'https': f'http://{host}:{port}'
            }
        return None

    def start_conversion(self):
        """开始转换过程"""
        input_path = self.input_path_var.get()
        output_path = self.output_path_var.get()
        if not input_path or not output_path:
            messagebox.showerror("错误", "请选择输入文件和输出目录")
            return
        if not os.path.exists(input_path):
            messagebox.showerror("错误", "输入文件不存在")
            return
        if not os.path.exists(output_path):
            try:
                os.makedirs(output_path)
            except Exception as e:
                messagebox.showerror("错误", f"创建输出目录失败: {str(e)}")
                return
        self.convert_button.configure(state='disabled')
        self.status_var.set("准备转换...")
        self.progress_var.set(0)
        try:
            proxy = self.get_current_proxy()
            if proxy:
                os.environ['HTTP_PROXY'] = proxy['http']
                os.environ['HTTPS_PROXY'] = proxy['https']
            else:
                os.environ.pop('HTTP_PROXY', None)
                os.environ.pop('HTTPS_PROXY', None)
            self.status_var.set("正在转换...")
            audio_paths = batch_text_to_speech(input_path, output_path, self.update_progress)

            if audio_paths:
                self.status_var.set(f"转换完成，共生成 {len(audio_paths)} 个音频文件")
                messagebox.showinfo("完成", f"成功生成 {len(audio_paths)} 个音频文件")
            else:
                self.status_var.set("转换完成，但未生成音频文件")
                messagebox.showwarning("警告", "未找到可转换的梵文内容")
        except Exception as e:
            self.status_var.set("转换失败")
            messagebox.showerror("错误", f"转换过程中发生错误: {str(e)}")
            logging.error(f"转换失败: {str(e)}")
        finally:
            self.convert_button.configure(state='normal')
            os.environ.pop('HTTP_PROXY', None)
            os.environ.pop('HTTPS_PROXY', None)

if __name__ == '__main__':
    logging.basicConfig(
        level=logging.DEBUG,  # 设置日志级别为 DEBUG
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.StreamHandler(),
            logging.FileHandler('audio_converter.log', encoding='utf-8')
        ]
    )
    try:
        app = SanskritAudioConverterGUI()
        app.mainloop()
    except Exception as e:
        logging.error(f"程序运行出错: {str(e)}")
        messagebox.showerror("错误", str(e))